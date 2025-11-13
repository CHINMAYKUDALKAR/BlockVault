"""DOCX redaction implementation using python-docx."""

from __future__ import annotations

import hashlib
from typing import Iterable, List

from docx import Document  # type: ignore

from .base_redactor import BaseRedactor, RedactionRequest, RedactionResult


class DOCXRedactor(BaseRedactor):
    """Performs deterministic text replacement and metadata scrubbing for DOCX files."""

    def apply_redactions(self) -> RedactionResult:
        doc = Document(self.request.source_path)

        removed_terms: List[str] = []
        terms = set(term for term in self.request.custom_terms if term)
        for pattern_term in self.request.patterns_applied:
            terms.add(pattern_term)
        for text in self.request.matched_texts:
            terms.add(text)

        if terms:
            removed_terms.extend(sorted(terms))
            for paragraph in doc.paragraphs:
                for term in terms:
                    if term in paragraph.text:
                        paragraph.text = paragraph.text.replace(term, "█" * len(term))

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for term in terms:
                            if term in cell.text:
                                cell.text = cell.text.replace(term, "█" * len(term))

        self._strip_metadata(doc)

        doc.save(self.request.output_path)
        sha256 = hashlib.sha256(self.request.output_path.read_bytes()).hexdigest()

        return RedactionResult(
            redacted_path=self.request.output_path,
            sha256=sha256,
            total_regions=len(self.request.regions),
            removed_terms=removed_terms,
        )

    def _strip_metadata(self, doc: Document) -> None:
        """Remove comments and sensitive metadata from the DOCX document."""

        core_props = doc.core_properties
        core_props.author = None
        core_props.comments = None
        core_props.keywords = None
        core_props.last_modified_by = self.request.requested_by
        core_props.revision = 1
        core_props.subject = None
        core_props.title = None


